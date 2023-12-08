import requests
from bs4 import BeautifulSoup
from docx import Document

url = 'https://dar-vetclinic.ru/encziklopediya/'  # Устанавливаем URL-адрес страницы для извлечения данных
response = requests.get(url)  # Отправляем GET-запрос на указанный URL-адрес и получаем ответ
soup = BeautifulSoup(response.text, 'html.parser')  # Создаем объект BeautifulSoup для парсинга HTML-кода страницы
links = soup.find_all('a')  # Ищем все ссылки на странице
visited_links = set()  # Создаем множество для отслеживания посещенных ссылок
doc = Document()  # Создаем пустой документ для хранения данных

for link in links:  # Проходимся по каждой ссылке на странице
    href = link.get('href')  # Получаем URL-адрес из ссылки
    if '/sobaki/' in href and href not in visited_links:  # Проверяем, что ссылка содержит "/sobaki/" и еще не была посещена
        visited_links.add(href)  # Добавляем ссылку в множество посещенных ссылок
        response = requests.get(href)  # Отправляем GET-запрос на ссылку и получаем ответ
        subpage_soup = BeautifulSoup(response.text, 'html.parser')  # Создаем объект BeautifulSoup для парсинга HTML-кода подстраницы
        entry_title = subpage_soup.find('h1', class_='entry-title')  # Ищем заголовок записи
        if entry_title:  # Проверяем, что заголовок найден
            entry_title_text = entry_title.get_text()  # Получаем текст заголовка
            article_body_container = subpage_soup.find('div', class_='article-body-container')  # Ищем контейнер с текстом статьи
            if article_body_container:  # Проверяем, что контейнер найден
                article_text = article_body_container.get_text(separator='\n')  # Получаем текст статьи
                article_text_list = article_text.split("\n")  # Разделяем текст статьи на список по символу новой строки
                if 'ДАР»' in article_text_list:  # Проверяем наличие определенной строки в списке
                    first_index = article_text_list.index('ДАР»')  # Находим индекс этой строки
                    trimmed_list = article_text_list[:first_index]  # Обрезаем список до указанного индекса
                    trimmed_list1 = trimmed_list[:-4]  # Дополнительное обрезание списка
                    table = subpage_soup.find('table', class_='prices-table__services')  # Ищем таблицу с ценами
                    headers = [header.text for header in table.find_all('td', class_='prices-table-title-sec')]  # Получаем заголовки таблицы
                    tabl = []  # Создаем пустой список для хранения данных таблицы
                    for row in table.find_all('tr'):  # Проходимся по каждой строке таблицы
                        cells = row.find_all('td')  # Ищем все ячейки в строке
                        if len(cells) == len(headers):  # Проверяем, что количество ячеек равно количеству заголовков
                            tabl.append([cell.text for cell in cells])  # Добавляем данные ячеек в список
                    doc.add_heading(entry_title_text, level=1)  # Добавляем заголовок в документ
                    doc.add_paragraph('\n'.join(trimmed_list1))  # Добавляем текст статьи в документ
                    table = doc.add_table(rows=len(tabl), cols=len(headers))  # Добавляем таблицу в документ
                    for i in range(len(tabl)):  # Проходимся по каждой строке таблицы
                        for j in range(len(headers)):  # Проходимся по каждой ячейке в строке
                            table.cell(i, j).text = tabl[i][j]  # Добавляем данные из списка в ячейки таблицы.

# Сохраняем документ DOCX
doc.save('df.docx')
